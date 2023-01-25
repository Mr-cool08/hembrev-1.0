import datetime   
from docx import Document
from docx.shared import Inches
import os
import win32com.client as win32
import configparser
dt = datetime.date.today()
wk = dt.isocalendar()[1]
document = Document()
olApp = win32.Dispatch('Outlook.Application')
olNS = olApp.GetNameSpace('MAPI')
mailItem = olApp.CreateItem(0)
mailItem.Subject = 'Dummy Email'
mailItem.BodyFormat = 1
mailItem.Body = "Hello World"
mailItem.To = 'liam@suorsa.se' 


config = configparser.ConfigParser()
config.read('configfile.ini')
mail1 = config.get('Emails', 'email1')
mail2 = config.get('Emails', 'email2')
mail3 = config.get('Emails', 'Email3')
mail4 = config.get('Emails', 'Email4')
mail5 = config.get('Emails', 'Email5')

document.add_heading('Hembrev', 0)

p = document.add_paragraph('Skolarbete')

armatte = input("vad gör du i matten? ")
lermatte = input("Vad har du lärt dig i matten? ")
notematte = input("Har du någon note för matten? ")

arsv = input("Vad gör du i Svenskan? ")
lersv = input("Vad har du lärt dig i svenskan? ")
notesv = input("har du någon note för svenskan? ")

areng = input("Vad gör du i Engelskan? ")
lereng = input("Vad har du lärt dig i engelskan? ")
noteeng = input("har du någon note för engelskan? ")

arno = input("Vad gör du i NO? ")
lerno = input("Vad har du lärt dig i NO? ")
noteno = input("har du någon note för NO? ")

arso = input("Vad gör du i SO? ")
lerso = input("Vad har du lärt dig i SO? ")
noteso = input("har du någon note för SO? ")

aridh = input("Vad gör du i idrott? ")
leridh = input("Vad har du lärt dig i Idrotten? ")
noteidh = input("har du någon note för idrotten? ")

armu = input("Vad gör du i musiken? ")
lermu = input("Vad har du lärt dig i musiken? ")
notemu = input("har du någon note för Musik? ")

arbi = input("Vad gör du i Bilden? ")
lerbi = input("Vad har du lärt dig i Bilden? ")
notebi = input("har du någon note för Bild? ")

arhkk = input("Vad gör du i hemkunskapen? ")
lerhkk = input("Vad har du lärt dig i hemkunskapen? ")
notehkk = input("har du någon note för hemkunskap? ")

arsp = input("Vad gör du i spanskan? ")
lersp = input("Vad har du lärt dig i spanskan? ")
notesp = input("har du någon note för Spanskan? ")

arsl = input("Vad gör du i slöjden? ")
lersl = input("Vad har du lärt dig i slöjden? ")
notesl = input("har du någon note för slöjden? ")
if notematte == "nej" or notematte == "no":
    notematte = "-"
if notesv == "nej" or notesv == "no":
    notesv = "-"
if noteeng == "nej" or noteeng == "no":
    noteeng = "-"
if noteno == "nej" or noteno == "no":
    noteno = "-"
if noteso == "nej" or noteso == "no":
    noteso = "-"
if noteidh == "nej" or noteidh == "no":
    noteidh = "-"
if notemu == "nej" or notemu == "no":
    notemu = "-"
if noteidh == "nej" or noteidh == "no":
    noteidh = "-"
if notebi == "nej" or notebi == "no":
    notebi = "-"
if notehkk == "nej" or notehkk == "no":
    notehkk = "-"
if notesl == "nej" or notesl == "no":
    notesl = "-"
if notesp == "nej" or notesp == "no":
    notesp = "-"


records = (
    ("Matte", armatte, lermatte, notematte),
    ("svenska", arsv, lersv, notesv),
    ("Engelska", areng, lereng, noteeng), 
    ("NO", arno, lerno, noteno),
    ("SO", arso, lerso, noteso),
    ("Idrott", aridh, leridh, noteidh),
    ("Musik", armu, lermu, notemu),
    ("Bild", arbi, lerbi, notebi),
    ("Hemkundskap", arhkk, lerhkk, notehkk),
    ("Slöjd", arsl, lersl, notesl),
    ("Spanska", arsp, lersp, notesp)
)

table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ämne:'
hdr_cells[1].text = 'Arbetsområde:'
hdr_cells[2].text = 'Detta har jag lärt mig:'
hdr_cells[3].text = 'Notes:'
for qty, id, desc, des in records:
    row_cells = table.add_row().cells
    row_cells[0].text = qty
    row_cells[1].text = id
    row_cells[2].text = desc
    row_cells[3].text = des
document.add_page_break()

document.save('hembrev v' + str(wk) + ".docx")
docname = 'hembrev v' + str(wk) + ".docx"
olApp = win32.Dispatch('Outlook.Application')
olNS = olApp.GetNameSpace('MAPI')

# construct the email item object
mailItem = olApp.CreateItem(0)
mailItem.Subject = 'Hembrev'
mailItem.BodyFormat = 1
mailItem.Body = "hembrev"
mailItem.To = mail1, mail2, mail3, mail4, mail5


mailItem.Attachments.Add(os.path.join(os.getcwd(), docname))

mailItem.Display()

mailItem.Save()
mailItem.Send()