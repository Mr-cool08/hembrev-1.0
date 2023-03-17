import tkinter as tk
from tkinter import ttk
from docx import Document
import datetime
import smtplib
import datetime
from tkinter import messagebox
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import configparser
from windtalker import SymmetricCipher
import os
import random
#getting the week number
dt = datetime.date.today()
wk = dt.isocalendar()[1]





    
    

class Application(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.grid()
        self.create_widgets()
    
    #mark these as empty if the user dont input    
    #mark these as empty if the user dont input    
    händelse_vad1 = ""
    händelse_när1 = ""
    händelse_vad2 = ""
    händelse_när2 = ""
    händelse_vad3 = ""
    händelse_när3 = ""
    rest_vad1 = ""
    rest_hur1 = ""
    rest_när1 = ""
    rest_vad2 = ""
    rest_hur2 = ""
    rest_när2 = ""
    rest_vad3 = ""
    rest_hur3 = ""
    rest_när3 = ""
    #when pressed tab next input
    #when pressed tab next input
    def focus_next_widget(self, event):
        event.widget.tk_focusNext().focus()
    
    
    #encrypt and decrypt for the password
    #encrypt and decrypt for the password
    def encrypt(self):
        c = SymmetricCipher(password="Super secret password")
        c = SymmetricCipher(password="Super secret password")
        try:
            c.encrypt_file("password.txt")
            os.remove("password.txt")
        except OSError:
            os.remove("password-encrypted.txt")
    def decrypt(self):
        c = SymmetricCipher(password="Super secret password")
        c = SymmetricCipher(password="Super secret password")
        try:
            c.decrypt_file("password-encrypted.txt")
            os.remove("password-encrypted.txt")
        except:
            pass
    

    #if the user want to send something in the mail
    #if the user want to send something in the mail
    def mail_message(self):
        self.title_label = tk.Label(self, text="Skriv här vad som ska stå i mailet")
        self.title_label.grid(row=23, column=1, pady=10, columnspan=2)
        frame = tk.Frame(self)
        frame.grid(row=24, column=1, sticky="w", columnspan=50)
        self.message = tk.Entry(frame, width=50)
        self.message.grid(row=0, column=0)
        self.message.pack()



        



    #if the user have any rester
    #if the user have any rester
    def rest(self):
            
            
        self.title_label = tk.Label(self, text="Vad")
        self.title_label.grid(row=14, column=0, columnspan=3, pady=10)
        self.title_label = tk.Label(self, text="Hur")
        self.title_label.grid(row=14, column=1, columnspan=5, pady=10)
        self.title_label = tk.Label(self, text="När")
        self.title_label.grid(row=14, column=3, columnspan=6, pady=10)
        self.extra_label1 = tk.Label(self, text="Rest 1")
        self.extra_label1.grid(row=15, column=0)

        self.rest_vad1 = tk.Entry(self)
        self.rest_vad1.grid(row=15, column=1)
        self.rest_vad1.bind("<Tab>", lambda event: self.rest_hur1.focus_set())

        self.rest_hur1 = tk.Entry(self)
        self.rest_hur1.grid(row=15, column=2)
        self.rest_hur1.bind("<Tab>", lambda event: self.rest_nä1.focus_set())

        self.rest_när1 = tk.Entry(self)
        self.rest_när1.grid(row=15, column=3)
        
        self.extra_label2 = tk.Label(self, text="Rest 2")
        self.extra_label2.grid(row=16, column=0)

        self.rest_vad2 = tk.Entry(self)
        self.rest_vad2.grid(row=16, column=1)
        self.rest_vad2.bind("<Tab>", lambda event: self.extra_lärt2.focus_set())

        self.rest_hur2 = tk.Entry(self)
        self.rest_hur2.grid(row=16, column=2)
        self.rest_hur2.bind("<Tab>", lambda event: self.extra_note2.focus_set())
        
        self.rest_när2 = tk.Entry(self)
        self.rest_när2.grid(row=16, column=3)
        
        self.extra_label3 = tk.Label(self, text="Rest 3")
        self.extra_label3.grid(row=17, column=0)

        self.rest_vad3 = tk.Entry(self)
        self.rest_vad3.grid(row=17, column=1)
        self.rest_vad3.bind("<Tab>", lambda event: self.extra_lärt2.focus_set())

        self.rest_hur3 = tk.Entry(self)
        self.rest_hur3.grid(row=17, column=2)
        self.rest_hur3.bind("<Tab>", lambda event: self.extra_note2.focus_set())

        self.rest_när3 = tk.Entry(self)
        self.rest_när3.grid(row=17, column=3)
        tk.Entry.pack(self)
        
        
    #if the user have anything that is going to happen    
    #if the user have anything that is going to happen    
    def händelser(self):
            
            
        self.title_label = tk.Label(self, text="Vad")
        self.title_label.grid(row=18, column=0, columnspan=3, pady=10)
        self.title_label = tk.Label(self, text="När")
        self.title_label.grid(row=18, column=1, columnspan=3, pady=10)
        
        self.extra_label1 = tk.Label(self, text="Händelse 1")
        self.extra_label1.grid(row=19, column=0)

        self.händelse_vad1 = tk.Entry(self)
        self.händelse_vad1.grid(row=19, column=1)
        self.händelse_vad1.bind("<Tab>", lambda event: self.rest_hur1.focus_set())

        self.händelse_när1 = tk.Entry(self)
        self.händelse_när1.grid(row=19, column=2)
        self.händelse_när1.bind("<Tab>", lambda event: self.rest_nä1.focus_set())

        
        self.extra_label2 = tk.Label(self, text="Händelse 2")
        self.extra_label2.grid(row=20, column=0)

        self.händelse_vad2 = tk.Entry(self)
        self.händelse_vad2.grid(row=20, column=1)
        self.händelse_vad2.bind("<Tab>", lambda event: self.extra_lärt2.focus_set())

        self.händelse_när2 = tk.Entry(self)
        self.händelse_när2.grid(row=20, column=2)
        self.händelse_när2.bind("<Tab>", lambda event: self.extra_note2.focus_set())
        
        
        self.extra_label3 = tk.Label(self, text="Händelse 3")
        self.extra_label3.grid(row=21, column=0)

        self.händelse_vad3 = tk.Entry(self)
        self.händelse_vad3.grid(row=21, column=1)
        self.händelse_vad3.bind("<Tab>", lambda event: self.extra_lärt2.focus_set())

        self.händelse_när3 = tk.Entry(self)
        self.händelse_när3.grid(row=21, column=2)
        self.händelse_när3.bind("<Tab>", lambda event: self.extra_note2.focus_set())
     #creating the document with all inputs
    def create_document(self):
        
        document = Document()

        document.add_heading('Hembrev', 0)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ämne'
        hdr_cells[1].text = 'ARBETSOMRÅDE'
        hdr_cells[2].text = 'DETTA HAR JAG LÄRT MIG:'
        hdr_cells[3].text = 'notes'

        armatte = self.matte_gör.get()
        lermatte = self.matte_lärt.get()
        notematte = self.matte_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Matte'
        row_cells[1].text = armatte
        row_cells[2].text = lermatte
        row_cells[3].text = notematte
        
        
        arsv = self.sv_gör.get()
        lersv = self.sv_lärt.get()
        notesv = self.sv_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Svenska'
        row_cells[1].text = arsv
        row_cells[2].text = lersv
        row_cells[3].text = notesv

        areng = self.eng_gör.get()
        lereng = self.eng_lärt.get()
        noteeng = self.eng_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Engelska'
        row_cells[1].text = areng
        row_cells[2].text = lereng
        row_cells[3].text = noteeng
        
        arno = self.no_gör.get()
        lerno = self.no_lärt.get()
        noteno = self.no_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'NO'
        row_cells[1].text = arno
        row_cells[2].text = lerno
        row_cells[3].text = noteno
        
        arso = self.so_gör.get()
        lerso = self.so_lärt.get()
        noteso = self.so_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'NO'
        row_cells[1].text = arso
        row_cells[2].text = lerso
        row_cells[3].text = noteso
        
        aridh = self.idh_gör.get()
        leridh = self.idh_lärt.get()
        noteidh = self.idh_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Idrott'
        row_cells[1].text = aridh
        row_cells[2].text = leridh
        row_cells[3].text = noteidh
        
        armu = self.mu_gör.get()
        lermu = self.mu_lärt.get()
        notemu = self.mu_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Musik'
        row_cells[1].text = armu
        row_cells[2].text = lermu
        row_cells[3].text = notemu
        
        arbi = self.bi_gör.get()
        lerbi = self.bi_lärt.get()
        notebi = self.bi_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Bild'
        row_cells[1].text = arbi
        row_cells[2].text = lerbi
        row_cells[3].text = notebi
        
        arhkk = self.hkk_gör.get()
        lerhkk = self.hkk_lärt.get()
        notehkk = self.hkk_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Hemkundskap'
        row_cells[1].text = arhkk
        row_cells[2].text = lerhkk
        row_cells[3].text = notehkk
        
        arsp = self.sp_gör.get()
        lersp = self.sp_lärt.get()
        notesp = self.sp_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Språk'
        row_cells[1].text = arsp
        row_cells[2].text = lersp
        row_cells[3].text = notesp
        
        arsl = self.sl_gör.get()
        lersl = self.sl_lärt.get()
        notesl = self.sl_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Slöjd'
        row_cells[1].text = arsl
        row_cells[2].text = lersl
        row_cells[3].text = notesl
        
    
    
        document.add_heading('Händelser', 0)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vad'
        hdr_cells[1].text = 'När'
        #only if the checkbox is checked it inputs it else it will just be empty
        #only if the checkbox is checked it inputs it else it will just be empty
        if self.händelser_checked.get():
            händelse_vad1 = self.händelse_vad1.get()
            händelse_när1 = self.händelse_när1.get() 
            händelse_när1 = self.händelse_när1.get() 
            row_cells = table.add_row().cells
            row_cells[0].text = händelse_vad1
            row_cells[1].text = händelse_när1
            
            händelse_vad2 = self.händelse_vad2.get()
            händelse_när2 = self.händelse_när2.get()
            row_cells = table.add_row().cells
            row_cells[0].text = händelse_vad2
            row_cells[1].text = händelse_när2

            händelse_vad3 = self.händelse_vad3.get()
            händelse_när3 = self.händelse_när3.get()
            row_cells = table.add_row().cells
            row_cells[0].text = händelse_vad3
            row_cells[1].text = händelse_när3

        
        document.add_heading('Rester', 0)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vad'
        hdr_cells[1].text = 'Hur'
        hdr_cells[2].text = 'När'
        if self.rest_checked.get():
            rest_vad1 = self.rest_vad1.get()
            rest_hur1 = self.rest_hur1.get()
            rest_när1 = self.rest_när1.get()
            row_cells = table.add_row().cells
            row_cells[0].text = rest_vad1
            row_cells[1].text = rest_hur1
            row_cells[2].text = rest_när1
            
            rest_vad2 = self.rest_vad2.get()
            rest_hur2 = self.rest_hur2.get()
            rest_när2 = self.rest_när2.get()
            row_cells = table.add_row().cells
            row_cells[0].text = rest_vad2
            row_cells[1].text = rest_hur2
            row_cells[2].text = rest_när2
            
            rest_vad3 = self.rest_vad3.get()
            rest_hur3 = self.rest_hur3.get()
            rest_när3 = self.rest_när3.get()
            row_cells = table.add_row().cells
            row_cells[0].text = rest_vad3
            row_cells[1].text = rest_hur3
            row_cells[2].text = rest_när3


        #saving the document with the week number
        #saving the document with the week number
        document.save(f'Hembrev v{wk}.docx')
        
    
    def sendmail(self, message):
            config = configparser.ConfigParser()
            config.read('config.ini')
            mail = config.get('login', 'email')
            name = mail.split(".")[0]
            good_bye = ["Ha en bra dag!:)", f"Mvh {name}", f"vänliga hälsningar {name}"]
            random_message = random.choice(good_bye)
            mail_body = message  + "\n" + "\n" +"\n" +"\n" +"\n" + random_message +"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"\n"+"Detta mail var skickat igenom hembrevs programmet"
            
            self.decrypt()
            config = configparser.ConfigParser()
            config.read('config.ini')
            mail1 = config.get('Emails', 'email1')
            mail2 = config.get('Emails', 'email2')
            mail3 = config.get('Emails', 'Email3')
            mail4 = config.get('Emails', 'Email4')
            mail5 = config.get('Emails', 'Email5')
            mail = config.get('login', 'email')
            with open('password.txt','r') as file:
                password = file.read()

            #making the email
            #making the email
            docname = 'hembrev v' + str(wk) + ".docx"
            msg = MIMEMultipart()
            msg['From'] = mail
            receivers = [mail1, mail2, mail3, mail4, mail5]
            msg['To'] = ', '.join(receivers)
            msg['Subject'] = 'hembrev v' + str(wk)
            body = mail_body
            msg.attach(MIMEText(body, 'plain'))
            attachment = open(docname, 'rb')
            part = MIMEBase('application', "octet-stream")
            part.set_payload((attachment).read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', "attachment; filename= %s" % docname)
            msg.attach(part)

            #sending the mail with a office 365 server
            #sending the mail with a office 365 server
            server = smtplib.SMTP('smtp.office365.com', 587)  ### put your relevant SMTP here
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(mail, password)  ### if applicable
            server.send_message(msg)
            server.quit()
            self.encrypt()
    #when the user presses the sumbit button
    #when the user presses the sumbit button
    def submit(self):
        if self.mail_message_checked.get() == 1:
            message = self.message.get().strip() or ' '
        else:
            message = ' '
        self.create_document()
        self.master.destroy()
        self.sendmail(message)



    
    #create_widgets is where all the visual stuff is
    def create_widgets(self):
        #labels so the user know where to input what
        #labels so the user know where to input what
        self.title_label = tk.Label(self, text="Vad gör du?")
        self.title_label.grid(row=0, column=1, pady=10, sticky="w")
        self.title_label = tk.Label(self, text="Vad lär du dig?")
        self.title_label.grid(row=0, column=2, columnspan=6, pady=10, sticky="w")
        self.title_label = tk.Label(self, text="Note")
        self.title_label.grid(row=0, column=3, columnspan=7, pady=10, sticky="w")


        #the input boxes for the subjects
        #the input boxes for the subjects
        def questions(self):
            self.matte_label = tk.Label(self, text="Matte")
            self.matte_label.grid(row=1, column=0, sticky="w")

            self.matte_gör = tk.Entry(self)
            self.matte_gör.grid(row=1, column=1, sticky="w")
            self.matte_gör.bind("<Tab>", lambda event: self.matte_lärt.focus_set())

            self.matte_lärt = tk.Entry(self)
            self.matte_lärt.grid(row=1, column=2, sticky="w")
            self.matte_lärt.bind("<Tab>", lambda event: self.matte_note.focus_set())

            self.matte_note = tk.Entry(self)
            self.matte_note.grid(row=1, column=3, sticky="w")
            self.matte_note.bind("<Tab>", lambda event: self.sv_gör.focus_set())

            self.sv_label = tk.Label(self, text="Svenska")
            self.sv_label.grid(row=2, column=0, sticky="w")

            self.sv_gör = tk.Entry(self)
            self.sv_gör.grid(row=2, column=1, sticky="w")
            self.sv_gör.bind("<Tab>", lambda event: self.sv_lärt.focus_set())

            self.sv_lärt = tk.Entry(self)
            self.sv_lärt.grid(row=2, column=2, sticky="w")
            self.sv_lärt.bind("<Tab>", lambda event: self.sv_note.focus_set())

            self.sv_note = tk.Entry(self)
            self.sv_note.grid(row=2, column=3, sticky="w")
            self.sv_note.bind("<Tab>", lambda event: self.eng_gör.focus_set())
            
            
            self.eng_label = tk.Label(self, text="Engleska")
            self.eng_label.grid(row=3, column=0, sticky="w")

            self.eng_gör = tk.Entry(self)
            self.eng_gör.grid(row=3, column=1)
            self.eng_gör.bind("<Tab>", lambda event: self.sv_lärt.focus_set())

            self.eng_lärt = tk.Entry(self)
            self.eng_lärt.grid(row=3, column=2)
            self.eng_lärt.bind("<Tab>", lambda event: self.sv_note.focus_set())

            self.eng_note = tk.Entry(self)
            self.eng_note.grid(row=3, column=3)
            self.eng_note.bind("<Tab>", lambda event: self.eng_gör.focus_set())
            
            self.no_label = tk.Label(self, text="NO")
            self.no_label.grid(row=4, column=0, sticky="w")

            self.no_gör = tk.Entry(self)
            self.no_gör.grid(row=4, column=1)
            self.no_gör.bind("<Tab>", lambda event: self.no_lärt.focus_set())

            self.no_lärt = tk.Entry(self)
            self.no_lärt.grid(row=4, column=2)
            self.no_lärt.bind("<Tab>", lambda event: self.no_note.focus_set())

            self.no_note = tk.Entry(self)
            self.no_note.grid(row=4, column=3)
            self.no_note.bind("<Tab>", lambda event: self.submit_button.focus_set())
            
            
            self.so_label = tk.Label(self, text="SO")
            self.so_label.grid(row=5, column=0, sticky="w")

            self.so_gör = tk.Entry(self)
            self.so_gör.grid(row=5, column=1)
            self.so_gör.bind("<Tab>", lambda event: self.so_lärt.focus_set())

            self.so_lärt = tk.Entry(self)
            self.so_lärt.grid(row=5, column=2)
            self.so_lärt.bind("<Tab>", lambda event: self.so_note.focus_set())

            self.so_note = tk.Entry(self)
            self.so_note.grid(row=5, column=3)
            self.so_note.bind("<Tab>", lambda event: self.idh_gör.focus_set())
            
            self.idh_label = tk.Label(self, text="IDH")
            self.idh_label.grid(row=6, column=0, sticky="w")

            self.idh_gör = tk.Entry(self)
            self.idh_gör.grid(row=6, column=1)
            self.idh_gör.bind("<Tab>", lambda event: self.idh_lärt.focus_set())

            self.idh_lärt = tk.Entry(self)
            self.idh_lärt.grid(row=6, column=2)
            self.idh_lärt.bind("<Tab>", lambda event: self.idh_note.focus_set())

            self.idh_note = tk.Entry(self)
            self.idh_note.grid(row=6, column=3)
            self.idh_note.bind("<Tab>", lambda event: self.mu_gör.focus_set())
            
            self.mu_label = tk.Label(self, text="Musik")
            self.mu_label.grid(row=7, column=0, sticky="w")

            self.mu_gör = tk.Entry(self)
            self.mu_gör.grid(row=7, column=1)
            self.mu_gör.bind("<Tab>", lambda event: self.mu_lärt.focus_set())

            self.mu_lärt = tk.Entry(self)
            self.mu_lärt.grid(row=7, column=2)
            self.mu_lärt.bind("<Tab>", lambda event: self.mu_note.focus_set())

            self.mu_note = tk.Entry(self)
            self.mu_note.grid(row=7, column=3)
            self.mu_note.bind("<Tab>", lambda event: self.bi_gör.focus_set())
            
            self.bi_label = tk.Label(self, text="Bild")
            self.bi_label.grid(row=8, column=0, sticky="w")

            self.bi_gör = tk.Entry(self)
            self.bi_gör.grid(row=8, column=1)
            self.bi_gör.bind("<Tab>", lambda event: self.bi_lärt.focus_set())

            self.bi_lärt = tk.Entry(self)
            self.bi_lärt.grid(row=8, column=2)
            self.bi_lärt.bind("<Tab>", lambda event: self.bi_note.focus_set())

            self.bi_note = tk.Entry(self)
            self.bi_note.grid(row=8, column=3)
            self.bi_note.bind("<Tab>", lambda event: self.hkk_gör.focus_set())
            
            self.hkk_label = tk.Label(self, text="hemkunskapen")
            self.hkk_label.grid(row=9, column=0, sticky="w")

            self.hkk_gör = tk.Entry(self)
            self.hkk_gör.grid(row=9, column=1)
            self.hkk_gör.bind("<Tab>", lambda event: self.hkk_lärt.focus_set())

            self.hkk_lärt = tk.Entry(self)
            self.hkk_lärt.grid(row=9, column=2)
            self.hkk_lärt.bind("<Tab>", lambda event: self.hkk_note.focus_set())
            

            self.hkk_note = tk.Entry(self)
            self.hkk_note.grid(row=9, column=3)
            self.hkk_note.bind("<Tab>", self.focus_next_widget)
            
            self.sp_label = tk.Label(self, text="Språk")
            self.sp_label.grid(row=10, column=0, sticky="w")

            self.sp_gör = tk.Entry(self)
            self.sp_gör.grid(row=10, column=1)
            self.sp_gör.bind("<Tab>", lambda event: self.sp_lärt.focus_set())

            self.sp_lärt = tk.Entry(self)
            self.sp_lärt.grid(row=10, column=2)
            self.sp_lärt.bind("<Tab>", lambda event: self.sp_note.focus_set())

            self.sp_note = tk.Entry(self)
            self.sp_note.grid(row=10, column=3)
            self.sp_note.bind("<Tab>", lambda event: self.ty_gör.focus_set())
            
            self.sl_label = tk.Label(self, text="Slöjd")
            self.sl_label.grid(row=11, column=0, sticky="w")

            self.sl_gör = tk.Entry(self)
            self.sl_gör.grid(row=11, column=1)
            self.sl_gör.bind("<Tab>", lambda event: self.sl_lärt.focus_set())

            self.sl_lärt = tk.Entry(self)
            self.sl_lärt.grid(row=11, column=2)
            self.sl_lärt.bind("<Tab>", lambda event: self.sl_note.focus_set())

            self.sl_note = tk.Entry(self)
            self.sl_note.grid(row=11, column=3)
            tk.Entry.pack(self)

        questions(self)
        

        #buttons and checkboxes
        #buttons and checkboxes
        self.rest_checked = tk.IntVar()
        self.add_fields_checkbox = tk.Checkbutton(self, text="Har du några rester?", variable=self.rest_checked, command=self.rest)
        self.add_fields_checkbox.grid(row=48, column=0)
        
        self.händelser_checked = tk.IntVar()
        self.add_fields_checkbox = tk.Checkbutton(self, text="Har du några händelser?", command=self.händelser, variable=self.händelser_checked)
        self.add_fields_checkbox.grid(row=49, column=0)
        
        self.mail_message_checked = tk.IntVar()
        self.message = tk.Checkbutton(self, text="Vill du skriva något i mailet?", command=self.mail_message, variable=self.mail_message_checked)
        self.message.grid(row=47, column=0)
        
        self.submit_button = tk.Button(self, text="Submit", command=self.submit)
        self.submit_button.grid(row=50, column=1, columnspan=3, pady=10)

# creating the window

# creating the window
root = tk.Tk()
root.state('zoomed')
#if user presses the X the program confirms it
#if user presses the X the program confirms it
def on_closing():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        root.destroy()

root.protocol("WM_DELETE_WINDOW", on_closing)
app = Application(master=root)
app.mainloop()