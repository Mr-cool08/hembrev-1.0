import tkinter as tk
from tkinter import ttk
from docx import Document
import datetime
import smtplib
import datetime  
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import configparser
dt = datetime.date.today()
wk = dt.isocalendar()[1]
from windtalker import SymmetricCipher
import os



    
    

class Application(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.grid()
        self.create_widgets()

    def create_widgets(self):
        self.title_label = tk.Label(self, text="Vad gör du?")
        self.title_label.grid(row=0, column=0, columnspan=3, pady=10)
        self.title_label = tk.Label(self, text="Vad lär du dig?")
        self.title_label.grid(row=0, column=1, columnspan=5, pady=10)
        self.title_label = tk.Label(self, text="Note")
        self.title_label.grid(row=0, column=3, columnspan=6, pady=10)

# ... the rest of your input field code

        
        self.matte_label = tk.Label(self, text="Matte")
        self.matte_label.grid(row=1, column=0)

        self.matte_gör = tk.Entry(self)
        self.matte_gör.grid(row=1, column=1)
        self.matte_gör.bind("<Tab>", lambda event: self.matte_lärt.focus_set())

        self.matte_lärt = tk.Entry(self)
        self.matte_lärt.grid(row=1, column=2)
        self.matte_lärt.bind("<Tab>", lambda event: self.matte_note.focus_set())

        self.matte_note = tk.Entry(self)
        self.matte_note.grid(row=1, column=3)
        self.matte_note.bind("<Tab>", lambda event: self.sv_gör.focus_set())

        self.sv_label = tk.Label(self, text="Svenska")
        self.sv_label.grid(row=2, column=0)

        self.sv_gör = tk.Entry(self)
        self.sv_gör.grid(row=2, column=1)
        self.sv_gör.bind("<Tab>", lambda event: self.sv_lärt.focus_set())

        self.sv_lärt = tk.Entry(self)
        self.sv_lärt.grid(row=2, column=2)
        self.sv_lärt.bind("<Tab>", lambda event: self.sv_note.focus_set())

        self.sv_note = tk.Entry(self)
        self.sv_note.grid(row=2, column=3)
        self.sv_note.bind("<Tab>", lambda event: self.eng_gör.focus_set())
        
        
        self.eng_label = tk.Label(self, text="Engleska")
        self.eng_label.grid(row=3, column=0)

        self.eng_gör = tk.Entry(self)
        self.eng_gör.grid(row=3, column=1)
        self.eng_gör.bind("<Tab>", lambda event: self.sv_lärt.focus_set())

        self.eng_lärt = tk.Entry(self)
        self.eng_lärt.grid(row=3, column=2)
        self.eng_lärt.bind("<Tab>", lambda event: self.sv_note.focus_set())

        self.eng_note = tk.Entry(self)
        self.eng_note.grid(row=3, column=3)
        self.eng_note.bind("<Tab>", lambda event: self.eng_gör.focus_set())
        
        self.no_label = tk.Label(self, text="NO")
        self.no_label.grid(row=4, column=0)

        self.no_gör = tk.Entry(self)
        self.no_gör.grid(row=4, column=1)
        self.no_gör.bind("<Tab>", lambda event: self.no_lärt.focus_set())

        self.no_lärt = tk.Entry(self)
        self.no_lärt.grid(row=4, column=2)
        self.no_lärt.bind("<Tab>", lambda event: self.no_note.focus_set())

        self.no_note = tk.Entry(self)
        self.no_note.grid(row=4, column=3)
        self.no_note.bind("<Tab>", lambda event: self.submit_button.focus_set())
        
        
        self.so_label = tk.Label(self, text="SO")
        self.so_label.grid(row=5, column=0)

        self.so_gör = tk.Entry(self)
        self.so_gör.grid(row=5, column=1)
        self.so_gör.bind("<Tab>", lambda event: self.so_lärt.focus_set())

        self.so_lärt = tk.Entry(self)
        self.so_lärt.grid(row=5, column=2)
        self.so_lärt.bind("<Tab>", lambda event: self.so_note.focus_set())

        self.so_note = tk.Entry(self)
        self.so_note.grid(row=5, column=3)
        self.so_note.bind("<Tab>", lambda event: self.idh_gör.focus_set())
        
        self.idh_label = tk.Label(self, text="IDH")
        self.idh_label.grid(row=6, column=0)

        self.idh_gör = tk.Entry(self)
        self.idh_gör.grid(row=6, column=1)
        self.idh_gör.bind("<Tab>", lambda event: self.idh_lärt.focus_set())

        self.idh_lärt = tk.Entry(self)
        self.idh_lärt.grid(row=6, column=2)
        self.idh_lärt.bind("<Tab>", lambda event: self.idh_note.focus_set())

        self.idh_note = tk.Entry(self)
        self.idh_note.grid(row=6, column=3)
        self.idh_note.bind("<Tab>", lambda event: self.mu_gör.focus_set())
        
        self.mu_label = tk.Label(self, text="Musik")
        self.mu_label.grid(row=7, column=0)

        self.mu_gör = tk.Entry(self)
        self.mu_gör.grid(row=7, column=1)
        self.mu_gör.bind("<Tab>", lambda event: self.mu_lärt.focus_set())

        self.mu_lärt = tk.Entry(self)
        self.mu_lärt.grid(row=7, column=2)
        self.mu_lärt.bind("<Tab>", lambda event: self.mu_note.focus_set())

        self.mu_note = tk.Entry(self)
        self.mu_note.grid(row=7, column=3)
        self.mu_note.bind("<Tab>", lambda event: self.bi_gör.focus_set())
        
        self.bi_label = tk.Label(self, text="Bild")
        self.bi_label.grid(row=8, column=0)

        self.bi_gör = tk.Entry(self)
        self.bi_gör.grid(row=8, column=1)
        self.bi_gör.bind("<Tab>", lambda event: self.bi_lärt.focus_set())

        self.bi_lärt = tk.Entry(self)
        self.bi_lärt.grid(row=8, column=2)
        self.bi_lärt.bind("<Tab>", lambda event: self.bi_note.focus_set())

        self.bi_note = tk.Entry(self)
        self.bi_note.grid(row=8, column=3)
        self.bi_note.bind("<Tab>", lambda event: self.hkk_gör.focus_set())
        
        self.hkk_label = tk.Label(self, text="hemkunskapen")
        self.hkk_label.grid(row=9, column=0)

        self.hkk_gör = tk.Entry(self)
        self.hkk_gör.grid(row=9, column=1)
        self.hkk_gör.bind("<Tab>", lambda event: self.hkk_lärt.focus_set())

        self.hkk_lärt = tk.Entry(self)
        self.hkk_lärt.grid(row=9, column=2)
        self.hkk_lärt.bind("<Tab>", lambda event: self.hkk_note.focus_set())

        self.hkk_note = tk.Entry(self)
        self.hkk_note.grid(row=9, column=3)
        self.hkk_note.bind("<Tab>", lambda event: self.sp_gör.focus_set())
        
        self.sp_label = tk.Label(self, text="Spanska")
        self.sp_label.grid(row=10, column=0)

        self.sp_gör = tk.Entry(self)
        self.sp_gör.grid(row=10, column=1)
        self.sp_gör.bind("<Tab>", lambda event: self.sp_lärt.focus_set())

        self.sp_lärt = tk.Entry(self)
        self.sp_lärt.grid(row=10, column=2)
        self.sp_lärt.bind("<Tab>", lambda event: self.sp_note.focus_set())

        self.sp_note = tk.Entry(self)
        self.sp_note.grid(row=10, column=3)
        self.sp_note.bind("<Tab>", lambda event: self.ty_gör.focus_set())
        
        self.ty_label = tk.Label(self, text="Tyska")
        self.ty_label.grid(row=11, column=0)

        self.ty_gör = tk.Entry(self)
        self.ty_gör.grid(row=11, column=1)
        self.ty_gör.bind("<Tab>", lambda event: self.ty_lärt.focus_set())

        self.ty_lärt = tk.Entry(self)
        self.ty_lärt.grid(row=11, column=2)
        self.ty_lärt.bind("<Tab>", lambda event: self.ty_note.focus_set())

        self.ty_note = tk.Entry(self)
        self.ty_note.grid(row=11, column=3)
        self.ty_note.bind("<Tab>", lambda event: self.sl_gör.focus_set())
        
        
        self.sl_label = tk.Label(self, text="Slöjd")
        self.sl_label.grid(row=12, column=0)

        self.sl_gör = tk.Entry(self)
        self.sl_gör.grid(row=12, column=1)
        self.sl_gör.bind("<Tab>", lambda event: self.sl_lärt.focus_set())

        self.sl_lärt = tk.Entry(self)
        self.sl_lärt.grid(row=12, column=2)
        self.sl_lärt.bind("<Tab>", lambda event: self.sl_note.focus_set())

        self.sl_note = tk.Entry(self)
        self.sl_note.grid(row=12, column=3)










        # repeat the process for all the subjects

        self.submit_button = tk.Button(self, text="Submit", command=self.submit)
        self.submit_button.grid(row=13, column=0, columnspan=4, pady=10)
    def encrypt():
        c = SymmetricCipher(password="password")
        try:
            c.encrypt_file("password.txt")
            os.remove("password.txt")
        except OSError:
            os.remove("password-encrypted.txt")
        
        
    def decrypt():
        c = SymmetricCipher(password="password")
        try:
            c.decrypt_file("password-encrypted.txt")
            os.remove("password-encrypted.txt")
        except:
            pass
    def submit(self):
        self.create_document()
        self.master.destroy()
        #self.sendmail()

    """def sendmail():
        decrypt()
        config = configparser.ConfigParser()
        config.read('config.ini')
        mail1 = config.get('Emails', 'email1')
        mail2 = config.get('Emails', 'email2')
        mail3 = config.get('Emails', 'Email3')
        mail4 = config.get('Emails', 'Email4')
        mail5 = config.get('Emails', 'Email5')
        mail = config.get('login', 'email')
        with open('password.txt','r') as file:
            password = file.read()


        docname = 'hembrev v' + str(wk) + ".docx"
        msg = MIMEMultipart()
        msg['From'] = mail
    ## vilka som ska få mailet
        receivers = [mail1, mail2, mail3, mail4, mail5]
        msg['To'] = ', '.join(receivers)
        msg['Subject'] = 'hembrev v' + str(wk)
        body = 'hembrev v' + str(wk)
        msg.attach(MIMEText(body, 'plain'))

    ## ATTACHMENT PART OF THE CODE IS HERE
        attachment = open(docname, 'rb')
        part = MIMEBase('application', "octet-stream")
        part.set_payload((attachment).read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', "attachment; filename= %s" % docname)
        msg.attach(part)

    # construct the email item object
        server = smtplib.SMTP('smtp.office365.com', 587)  ### put your relevant SMTP here
        server.ehlo()
        server.starttls()
        server.ehlo()
        server.login(mail, password)  ### if applicable
        server.send_message(msg)
        server.quit()
        encrypt()"""


    def create_document(self):
        dt = datetime.date.today()
        wk = dt.isocalendar()[1]
        document = Document()

        document.add_heading('Hembrev', 0)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ämne'
        hdr_cells[1].text = 'ARBETSOMRÅDE'
        hdr_cells[2].text = 'DETTA HAR JAG LÄRT MIG:'
        hdr_cells[3].text = 'notes'

        armatte = self.matte_gör.get()
        lermatte = self.matte_lärt.get()
        notematte = self.matte_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Matte'
        row_cells[1].text = armatte
        row_cells[2].text = lermatte
        row_cells[3].text = notematte
        arsv = self.sv_gör.get()
        lersv = self.sv_lärt.get()
        notesv = self.sv_note.get()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Svenska'
        row_cells[1].text = arsv
        row_cells[2].text = lersv
        row_cells[3].text = notesv

        areng = self.eng_gör.get()
        # repeat the process for all the subject

        document.save(f'Hembrev_vecka{wk}.docx')


        
        #get the value from all the entry fields
        
        # rest of the code 

root = tk.Tk()
app = Application(master=root)
app.mainloop()
